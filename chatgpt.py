import docx
import openai
from openai_api_key import api_key

# Initialize OpenAI API
openai.api_key = api_key

def extract_questions(input_doc):
    """Extracts questions from a Word document."""
    doc = docx.Document(input_doc)
    questions = []
    for para in doc.paragraphs:
        if para.text.endswith('?'):
            questions.append(para.text)
    return questions

def generate_answer(question):
    """Generates an answer using the ChatGPT API."""
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": question},
        ]
    )
    return response['choices'][0]['message']['content']

def insert_answers(input_doc, output_doc):
    """Inserts answers after each question in the Word document."""
    doc = docx.Document(input_doc)
    new_doc = docx.Document()

    for para in doc.paragraphs:
        new_doc.add_paragraph(para.text)
        if para.text.endswith('?'):
            answer = generate_answer(para.text)
            new_doc.add_paragraph(answer)  # Insert answer after the question

    new_doc.save(output_doc)

if __name__ == "__main__":
    input_doc = "input.docx"
    output_doc = "output.docx"
    
    insert_answers(input_doc, output_doc)
    print(f"Answers inserted and saved to {output_doc}")
